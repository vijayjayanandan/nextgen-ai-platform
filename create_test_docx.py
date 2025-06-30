"""
Create a test DOCX document for format validation testing.
"""

try:
    from docx import Document
    from docx.shared import Inches
    
    def create_eligibility_docx():
        """Create a DOCX document about eligibility criteria"""
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Eligibility Criteria for Government Services', 0)
        
        # Add content
        doc.add_heading('Primary Requirements', level=1)
        
        p1 = doc.add_paragraph('To be eligible for government services, applicants must meet the following criteria:')
        
        # Add numbered list
        doc.add_paragraph('Residency Status: Must be a legal resident or citizen of Canada', style='List Number')
        doc.add_paragraph('Age Requirements: Must be at least 18 years of age', style='List Number')
        doc.add_paragraph('Documentation: Must provide valid government-issued identification', style='List Number')
        doc.add_paragraph('Income Verification: Must meet minimum income thresholds as specified', style='List Number')
        
        doc.add_heading('Additional Considerations', level=1)
        
        considerations = [
            'Previous service history may be considered',
            'Special provisions exist for vulnerable populations',
            'Emergency exceptions may apply in certain circumstances'
        ]
        
        for consideration in considerations:
            doc.add_paragraph(consideration, style='List Bullet')
        
        doc.add_heading('Application Process', level=1)
        
        process_steps = [
            'Complete the online application form',
            'Submit required documentation',
            'Attend verification appointment if required',
            'Await processing and decision'
        ]
        
        for step in process_steps:
            doc.add_paragraph(step, style='List Number')
        
        # Add contact information
        doc.add_heading('Contact Information', level=1)
        doc.add_paragraph('For more information, contact the service center at 1-800-XXX-XXXX.')
        
        # Save the document
        filename = "test_documents/eligibility_criteria.docx"
        doc.save(filename)
        print(f"✅ Created DOCX: {filename}")
        return filename
    
    if __name__ == "__main__":
        create_eligibility_docx()
        
except ImportError:
    print("❌ python-docx not available. Creating RTF file as alternative...")
    
    # Create RTF file as fallback (can be read by Word)
    rtf_content = r"""{{\rtf1\ansi\deff0 {{\fonttbl {{\f0 Times New Roman;}}}}
\f0\fs24 
{\b\fs28 Eligibility Criteria for Government Services\par}
\par
{\b\fs24 Primary Requirements\par}
\par
To be eligible for government services, applicants must meet the following criteria:\par
\par
1. Residency Status: Must be a legal resident or citizen of Canada\par
2. Age Requirements: Must be at least 18 years of age\par
3. Documentation: Must provide valid government-issued identification\par
4. Income Verification: Must meet minimum income thresholds as specified\par
\par
{\b\fs24 Additional Considerations\par}
\par
• Previous service history may be considered\par
• Special provisions exist for vulnerable populations\par
• Emergency exceptions may apply in certain circumstances\par
\par
{\b\fs24 Application Process\par}
\par
1. Complete the online application form\par
2. Submit required documentation\par
3. Attend verification appointment if required\par
4. Await processing and decision\par
\par
{\b\fs24 Contact Information\par}
\par
For more information, contact the service center at 1-800-XXX-XXXX.\par
}}"""
    
    try:
        with open("test_documents/eligibility_criteria.rtf", "w", encoding='utf-8') as f:
            f.write(rtf_content)
        print("✅ Created RTF file: test_documents/eligibility_criteria.rtf")
    except Exception as e:
        print(f"❌ Failed to create RTF: {e}")
        
        # Final fallback - create a simple text file with .docx extension
        # (This won't be a real DOCX but will test the filename matching)
        simple_content = """Eligibility Criteria for Government Services

Primary Requirements:
To be eligible for government services, applicants must meet the following criteria:

1. Residency Status: Must be a legal resident or citizen of Canada
2. Age Requirements: Must be at least 18 years of age  
3. Documentation: Must provide valid government-issued identification
4. Income Verification: Must meet minimum income thresholds as specified

Additional Considerations:
- Previous service history may be considered
- Special provisions exist for vulnerable populations
- Emergency exceptions may apply in certain circumstances

Application Process:
1. Complete the online application form
2. Submit required documentation
3. Attend verification appointment if required
4. Await processing and decision

Contact Information:
For more information, contact the service center at 1-800-XXX-XXXX.
"""
        
        with open("test_documents/eligibility_criteria_text.docx", "w", encoding='utf-8') as f:
            f.write(simple_content)
        print("✅ Created text file with .docx extension: test_documents/eligibility_criteria_text.docx")
