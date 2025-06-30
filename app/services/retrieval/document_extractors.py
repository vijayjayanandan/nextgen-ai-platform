"""
Enterprise-grade document extraction system for diverse file formats.

This module provides a modular, extensible architecture for extracting text content
from various document formats including PDF, DOCX, HTML, Markdown, and plain text.

The system uses MIME-type detection and file extension analysis to route documents
to the appropriate extractor, ensuring robust handling of diverse file types.

Author: Senior Python Architect
Purpose: Government AI Platform Document Processing
"""

import re
import unicodedata
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Type, Optional
import logging

# Configure logging
logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    """Raised when a file type is not supported by any available extractor."""
    pass


class DocumentExtractionError(Exception):
    """Raised when document extraction fails due to corruption or other issues."""
    pass


class DocumentExtractor(ABC):
    """
    Abstract base class for document text extractors.
    
    All concrete extractors must implement the extract method to return
    clean, normalized UTF-8 plain text from document bytes.
    """
    
    @abstractmethod
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract clean text content from document bytes.
        
        Args:
            file_bytes: Raw bytes of the document file
            
        Returns:
            Clean, normalized UTF-8 plain text content
            
        Raises:
            DocumentExtractionError: If extraction fails
        """
        pass
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize text content for consistent processing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Normalized UTF-8 text with cleaned whitespace
        """
        if not text:
            return ""
        
        try:
            # Ensure UTF-8 encoding
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='replace')
            
            # Normalize unicode characters
            text = unicodedata.normalize('NFKC', text)
            
            # Clean up whitespace and line endings
            text = re.sub(r'\r\n|\r', '\n', text)  # Normalize line endings
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
            text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
            text = re.sub(r'[ \t]*\n[ \t]*', '\n', text)  # Clean line breaks
            
            # Remove control characters except newlines and tabs
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
            
            # Strip leading/trailing whitespace and ensure single trailing newline
            text = text.strip()
            if text:
                text += '\n'
            
            return text
            
        except Exception as e:
            logger.warning(f"Error normalizing text content: {e}")
            return text  # Return original if normalization fails


class TxtExtractor(DocumentExtractor):
    """Extractor for plain text files (.txt)."""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from plain text file bytes.
        
        Args:
            file_bytes: Raw bytes of the text file
            
        Returns:
            Normalized UTF-8 text content
            
        Raises:
            DocumentExtractionError: If text decoding fails
        """
        try:
            # Try multiple encodings to handle various text files
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    logger.debug(f"Successfully decoded text file with {encoding} encoding")
                    return self._normalize_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error replacement
            text = file_bytes.decode('utf-8', errors='replace')
            logger.warning("Text file decoded with error replacement")
            return self._normalize_text(text)
            
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract text from file: {e}")


class PdfExtractor(DocumentExtractor):
    """Extractor for PDF files (.pdf) using PyMuPDF."""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF file bytes using PyMuPDF.
        
        Args:
            file_bytes: Raw bytes of the PDF file
            
        Returns:
            Normalized UTF-8 text content
            
        Raises:
            DocumentExtractionError: If PDF extraction fails
        """
        try:
            import fitz  # PyMuPDF
            
            # Open PDF from bytes
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            # Check if PDF is encrypted
            if doc.needs_pass:
                doc.close()
                raise DocumentExtractionError("PDF is password protected")
            
            # Extract text from all pages
            text_content = []
            
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    
                    # Skip pages with no text (likely image-only)
                    if page_text.strip():
                        text_content.append(page_text)
                    else:
                        logger.debug(f"Page {page_num + 1} contains no extractable text")
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            doc.close()
            
            if not text_content:
                raise DocumentExtractionError("No text content found in PDF")
            
            # Join all pages with double newlines
            full_text = "\n\n".join(text_content)
            return self._normalize_text(full_text)
            
        except ImportError:
            raise DocumentExtractionError("PyMuPDF (fitz) not available for PDF extraction")
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract text from PDF: {e}")


class DocxExtractor(DocumentExtractor):
    """Extractor for Microsoft Word documents (.docx) using python-docx."""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX file bytes using python-docx.
        
        Args:
            file_bytes: Raw bytes of the DOCX file
            
        Returns:
            Normalized UTF-8 text content
            
        Raises:
            DocumentExtractionError: If DOCX extraction fails
        """
        try:
            from docx import Document
            from io import BytesIO
            
            # Open DOCX from bytes
            doc_stream = BytesIO(file_bytes)
            doc = Document(doc_stream)
            
            # Extract text from all paragraphs
            text_content = []
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_content.append(para_text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise DocumentExtractionError("No text content found in DOCX")
            
            # Join all content with newlines
            full_text = "\n".join(text_content)
            return self._normalize_text(full_text)
            
        except ImportError:
            raise DocumentExtractionError("python-docx not available for DOCX extraction")
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract text from DOCX: {e}")


class HtmlExtractor(DocumentExtractor):
    """Extractor for HTML files (.html, .htm) using BeautifulSoup."""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from HTML file bytes using BeautifulSoup.
        
        Args:
            file_bytes: Raw bytes of the HTML file
            
        Returns:
            Normalized UTF-8 text content with HTML tags removed
            
        Raises:
            DocumentExtractionError: If HTML extraction fails
        """
        try:
            from bs4 import BeautifulSoup
            
            # Decode HTML bytes
            try:
                html_content = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                # Try other common encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        html_content = file_bytes.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    html_content = file_bytes.decode('utf-8', errors='replace')
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "meta", "link", "noscript"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            if not text.strip():
                raise DocumentExtractionError("No text content found in HTML")
            
            return self._normalize_text(text)
            
        except ImportError:
            raise DocumentExtractionError("BeautifulSoup (bs4) not available for HTML extraction")
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract text from HTML: {e}")


class MarkdownExtractor(DocumentExtractor):
    """Extractor for Markdown files (.md, .markdown) with optional rendering."""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from Markdown file bytes.
        
        Attempts to render Markdown to HTML then extract text, with fallback
        to raw text if markdown library is not available.
        
        Args:
            file_bytes: Raw bytes of the Markdown file
            
        Returns:
            Normalized UTF-8 text content
            
        Raises:
            DocumentExtractionError: If Markdown extraction fails
        """
        try:
            # Decode Markdown bytes
            try:
                md_content = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                md_content = file_bytes.decode('utf-8', errors='replace')
            
            # Try to render Markdown to HTML then extract text
            try:
                import markdown
                from bs4 import BeautifulSoup
                
                # Convert Markdown to HTML
                html = markdown.markdown(md_content)
                
                # Extract text from HTML
                soup = BeautifulSoup(html, 'html.parser')
                text = soup.get_text()
                
                logger.debug("Markdown rendered to HTML and text extracted")
                
            except ImportError:
                # Fallback to raw text with basic Markdown cleanup
                logger.debug("Markdown library not available, using raw text extraction")
                text = self._clean_markdown_syntax(md_content)
            
            if not text.strip():
                raise DocumentExtractionError("No text content found in Markdown")
            
            return self._normalize_text(text)
            
        except Exception as e:
            raise DocumentExtractionError(f"Failed to extract text from Markdown: {e}")
    
    def _clean_markdown_syntax(self, md_text: str) -> str:
        """
        Clean basic Markdown syntax from raw text.
        
        Args:
            md_text: Raw Markdown text
            
        Returns:
            Text with basic Markdown syntax removed
        """
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', md_text, flags=re.MULTILINE)
        
        # Remove bold/italic markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Italic
        text = re.sub(r'__([^_]+)__', r'\1', text)      # Bold
        text = re.sub(r'_([^_]+)_', r'\1', text)        # Italic
        
        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Remove inline code
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
        
        # Remove horizontal rules
        text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
        
        # Remove list markers
        text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
        
        return text


def _detect_mime_type(file_bytes: bytes) -> Optional[str]:
    """
    Detect MIME type of file bytes using python-magic.
    
    Args:
        file_bytes: Raw file bytes
        
    Returns:
        MIME type string or None if detection fails
    """
    try:
        import magic
        
        # Try to detect MIME type
        mime_type = magic.from_buffer(file_bytes, mime=True)
        logger.debug(f"Detected MIME type: {mime_type}")
        return mime_type
        
    except ImportError:
        logger.warning("python-magic not available for MIME type detection")
        return None
    except Exception as e:
        logger.warning(f"MIME type detection failed: {e}")
        return None


def _get_file_extension(file_name: str) -> str:
    """
    Extract file extension from filename.
    
    Args:
        file_name: Name of the file
        
    Returns:
        Lowercase file extension without the dot
    """
    return Path(file_name).suffix.lower().lstrip('.')


def get_extractor(file_name: str, file_bytes: bytes) -> DocumentExtractor:
    """
    Factory method to get the appropriate document extractor.
    
    Uses both file extension and MIME type detection to determine
    the correct extractor for the given file.
    
    Args:
        file_name: Name of the file (used for extension detection)
        file_bytes: Raw bytes of the file (used for MIME detection)
        
    Returns:
        Appropriate DocumentExtractor instance
        
    Raises:
        UnsupportedFileTypeError: If file type is not supported
    """
    # File extension to extractor mapping
    EXTENSION_EXTRACTORS: Dict[str, Type[DocumentExtractor]] = {
        'txt': TxtExtractor,
        'pdf': PdfExtractor,
        'docx': DocxExtractor,
        'html': HtmlExtractor,
        'htm': HtmlExtractor,
        'md': MarkdownExtractor,
        'markdown': MarkdownExtractor,
    }
    
    # MIME type to extractor mapping
    MIME_EXTRACTORS: Dict[str, Type[DocumentExtractor]] = {
        'text/plain': TxtExtractor,
        'application/pdf': PdfExtractor,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocxExtractor,
        'text/html': HtmlExtractor,
        'application/xhtml+xml': HtmlExtractor,
        'text/markdown': MarkdownExtractor,
        'text/x-markdown': MarkdownExtractor,
    }
    
    # Get file extension
    file_extension = _get_file_extension(file_name)
    logger.debug(f"File extension detected: {file_extension}")
    
    # Detect MIME type
    mime_type = _detect_mime_type(file_bytes)
    
    # Determine extractor based on MIME type first (more reliable)
    extractor_class = None
    
    if mime_type and mime_type in MIME_EXTRACTORS:
        extractor_class = MIME_EXTRACTORS[mime_type]
        logger.debug(f"Extractor selected by MIME type: {extractor_class.__name__}")
    
    # Fallback to file extension if MIME detection failed or didn't match
    elif file_extension in EXTENSION_EXTRACTORS:
        extractor_class = EXTENSION_EXTRACTORS[file_extension]
        logger.debug(f"Extractor selected by file extension: {extractor_class.__name__}")
    
    # Special handling for common MIME type variations
    elif mime_type:
        if mime_type.startswith('text/'):
            # Generic text files
            extractor_class = TxtExtractor
            logger.debug("Generic text extractor selected for text/* MIME type")
        elif 'html' in mime_type.lower():
            # HTML variants
            extractor_class = HtmlExtractor
            logger.debug("HTML extractor selected for HTML-like MIME type")
    
    if extractor_class is None:
        # Create detailed error message
        error_msg = f"Unsupported file type: '{file_name}'"
        if file_extension:
            error_msg += f" (extension: .{file_extension})"
        if mime_type:
            error_msg += f" (MIME: {mime_type})"
        
        supported_extensions = list(EXTENSION_EXTRACTORS.keys())
        error_msg += f". Supported extensions: {', '.join(supported_extensions)}"
        
        raise UnsupportedFileTypeError(error_msg)
    
    # Log the selected extractor for debugging
    logger.debug("Extractor selected: %s for file: %s", extractor_class.__name__, file_name)
    
    # Return instance of the selected extractor
    return extractor_class()


# Export public interface
__all__ = [
    'DocumentExtractor',
    'TxtExtractor',
    'PdfExtractor', 
    'DocxExtractor',
    'HtmlExtractor',
    'MarkdownExtractor',
    'get_extractor',
    'UnsupportedFileTypeError',
    'DocumentExtractionError',
]
